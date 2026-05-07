#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Edit the graduation template in-place as a styled working copy.

This script keeps the original Word template as the base document, preserves its
styles, cover/declaration pages, and replaces the sample thesis body from
"摘 要" onward with the generated graduation report content.
"""
from __future__ import annotations

import re
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = next(ROOT.glob("4*.docx"))
MARKDOWN = ROOT / "docs" / "graduation_report.md"
OUTPUT = ROOT / "刘沛林_毕业设计报告_模板编辑版.docx"


def replace_visible_text(doc: Document) -> None:
    """Replace placeholders visible to python-docx while preserving runs."""
    if len(doc.tables) > 1:
        info = {
            "院    系": "计算机科学与技术",
            "专业班级": "CS2207",
            "姓    名": "刘沛林",
            "学    号": "待补充",
            "指导教师": "待补充",
        }
        for row in doc.tables[1].rows:
            key = row.cells[0].text.strip()
            if key in info:
                row.cells[1].text = info[key]

    replace_map = {
        "XXX系统的设计与实现": "vLLM 在线推理服务性能评测与优化系统的设计与实现",
        "小岳岳": "刘沛林",
        "计科2201": "CS2207",
        "U202215102": "待补充",
        "郭德纲": "待补充",
        "2026年5月14日": "2026年5月",
    }

    def replace_in_paragraph(paragraph) -> None:
        for old, new in replace_map.items():
            if old in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(old, new)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def clear_from_abstract(doc: Document) -> None:
    """Remove all sample body elements from the template's abstract onward."""
    start_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.replace(" ", "").strip() == "摘要":
            start_idx = idx
            break
    if start_idx is None:
        raise RuntimeError("未找到模板中的“摘 要”锚点")

    body = doc._body._element
    start_el = doc.paragraphs[start_idx]._element
    children = list(body)
    start_pos = children.index(start_el)
    sect_pr = children[-1] if children and children[-1].tag.endswith("sectPr") else None
    for element in list(body)[start_pos:]:
        if sect_pr is not None and element is sect_pr:
            continue
        body.remove(element)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def render_markdown_body(doc: Document) -> None:
    """Render the report markdown using styles already defined in the template."""
    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    first = next((i for i, line in enumerate(lines) if line.strip() == "## 摘 要"), 0)
    lines = lines[first:]

    idx = 0
    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped:
            idx += 1
            continue

        if (
            stripped.startswith("|")
            and idx + 1 < len(lines)
            and lines[idx + 1].strip().startswith("|")
            and "---" in lines[idx + 1]
        ):
            rows: list[list[str]] = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                parts = [
                    cell.strip().replace("`", "")
                    for cell in lines[idx].strip().strip("|").split("|")
                ]
                if not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in parts):
                    rows.append(parts)
                idx += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
                for row_idx, row in enumerate(rows):
                    for col_idx, value in enumerate(row):
                        table.cell(row_idx, col_idx).text = value
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip().replace("摘 要", "摘  要")
            if text in ("摘  要", "Abstract"):
                add_paragraph(doc, text, style="Title", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
            else:
                style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Normal")
                add_paragraph(doc, text, style=style)
            idx += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            add_paragraph(doc, numbered.group(2).replace("`", ""), style="List Paragraph")
            idx += 1
            continue

        if stripped.startswith("- "):
            add_paragraph(doc, stripped[2:].replace("`", ""), style="List Paragraph")
            idx += 1
            continue

        if (
            stripped.startswith("**")
            and stripped.endswith("**")
            and "关键词" not in stripped
            and "Keywords" not in stripped
        ):
            idx += 1
            continue

        add_paragraph(doc, stripped.replace("`", ""), style="Normal")
        idx += 1


def insert_toc_before_first_chapter(doc: Document) -> None:
    first_h1 = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "绪论":
            first_h1 = paragraph
            break
    if first_h1 is None:
        return

    title = first_h1.insert_paragraph_before()
    title.style = doc.styles["Title"]
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.add_run("目  录")

    toc = first_h1.insert_paragraph_before()
    run = toc.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "请在 Word 中右键更新域生成目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for element in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(element)

    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    copyfile(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))
    replace_visible_text(doc)
    clear_from_abstract(doc)
    render_markdown_body(doc)
    insert_toc_before_first_chapter(doc)
    doc.save(str(OUTPUT))
    print(f"已生成模板编辑版: {OUTPUT} ({OUTPUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
